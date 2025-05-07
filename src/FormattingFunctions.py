# formatting functions
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt
from docx import Document

# Formatting Variables
universalFont = "Times New Roman"
universalColor = RGBColor(0, 0, 0)
cellColor = "2B8DBF"

def format_text(run, font_size, color=universalColor, bold_status=False, underline_status=False, italicize=False):
    """
    formats the text contained in run
    Inputs:
        - run: the run field which contains your text
        - font_size: font size for the text,
        - color: color in RGB, default is black RGBColor(0, 0, 0)
        - bold_status: whether the text is bold, default False
        - underline_status: whether the text is underlined, default False
        - italicize: whether the text is italic, default False
    """
    run.bold = bold_status
    run.underline = underline_status
    run.font.name = universalFont
    run.font.size = Pt(font_size)
    run.font.color.rgb = color
    run.italic = italicize

def format_table_width(cell, width):
    """
    Set width of a cell
    """
    width = int(width * 1440)

    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    
    for child in tcPr.findall(qn('w:tcW')):
        tcPr.remove(child)
    
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width))
    tcW.set(qn('w:type'), "dxa")
    tcPr.append(tcW)
    
def format_table_height(row, height):
    """ Set height of a table row in inches using XML properties """
    height = int(height * 1440)  # Convert inches to twips
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(ns.qn('w:val'), str(height))
    trHeight.set(ns.qn('w:hRule'), "exact")  # Use "exact" to force a fixed height
    row._element.get_or_add_trPr().append(trHeight)

def format_summary_page(result_table):
    """
    format the table on the summary page
    """
    for cell in result_table.rows[0].cells:
        format_cell_color(cell, cellColor)
    
    for row in result_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    format_text(run, 12)

def format_cell_color(cell, color):
    """
    fills the cell with the provided color.
    color must be in hex
    """
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def format_table_text(table, font_size, color=universalColor, bold=False, underline=False, italics=False):
    """
    formats the text of the cell in the provided format
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    format_text(run, font_size, color, bold, underline, italics)

def format_cell_padding(cell, top=0, right=0, bottom=0, left=0):
    """
    formats the gap between the table cell text and table border
    """
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    for side, value in (('top', top), ('start', left), ('bottom', bottom), ('end', right)):
        mar = OxmlElement(f'w:{side}')
        mar.set(qn('w:w'), str(value))
        mar.set(qn('w:type'), 'dxa')
        tcMar.append(mar)
    
    tcPr.append(tcMar)

def format_paragraph_spacing(paragraph, before=0, after=0):
    """
    formats the spacing before and after a paragraph
    """
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)

def add_field(paragraph):
    """
    Adds a date time field
    """
    run = paragraph.add_run()

    fldChar1 = OxmlElement('w:fldChar')  # Begin field
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')  # Field instruction
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'DATE \\@ "dd/MM/yyyy"'

    fldChar2 = OxmlElement('w:fldChar')  # Separate
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')  # End field
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def add_page_number(run, field):
    """
    add a field, generally page number, to the run
    Inputs:
        - run: the run which we add to
        - field: the type of field to add, usually PAGE or NUMPAGES
    """

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = field

    fldChar2 = OxmlElement('w:fldChar')  # Separate
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')  # End field
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

# TODO 
def insert_line_break(cur_paragraph):
    """
    insert a line "break" at the current location in the document
    """
    p = cur_paragraph._p
    pPr = p.get_or_add_pPr()

    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')       # Line style
    bottom.set(qn('w:sz'), '12')            # Thickness
    bottom.set(qn('w:space'), '1')          # Space between line and text
    bottom.set(qn('w:color'), '000000')     # Color in hex, 000000=black

    pBdr.append(bottom)
    pPr.append(pBdr)