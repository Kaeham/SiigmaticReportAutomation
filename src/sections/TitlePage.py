from src.utils.Formatter import *
from src.utils.AssetPath import get_asset_path
from src.utils.Constants import CIILOCK_LOGO, NATA_LOGO, SIGNATURE_IMAGE
from docx.shared import RGBColor

def write(doc, data):
    _add_header(doc)
    _add_body(doc, data.body_info)
    _add_signature(doc)
    _add_footer(doc, data.body_info[2])  # Assuming Report Number is at index 2
    doc.add_page_break()

def _add_header(doc):
    section = doc.sections[0]
    header = section.header

    table = header.add_table(rows=1, cols=3, width=section.page_width+section.right_margin)
    table.autofit = False
    table.cell(0, 0).paragraphs[0].add_run().add_picture(get_asset_path(CIILOCK_LOGO))
    table.cell(0, 2).paragraphs[0].add_run().add_picture(get_asset_path(NATA_LOGO))
    
    p = table.cell(0, 1).paragraphs[0]
    format_text(p.add_run("CiiLOCK Engineering\n"), 20, color=RGBColor(54, 95, 145))
    format_text(p.add_run("ABN: 90 112 576 815     PH.\n"), 12)
    format_text(p.add_run("18 Technology Circuit Factory 2, Hallam Vic 3803"), 12)
    table.cell(0, 2).paragraphs[0].alignment = 2
    format_table_width(table.cell(0, 0), 1)
    format_table_width(table.cell(0, 1), 4)
    format_table_width(table.cell(0, 2), 2)
    format_table_height(table.rows[0], 1)
    section.header_distance = Pt(0)


def _add_body(doc, body_info):
    format_text(doc.add_paragraph().add_run(""), 20)
    format_text(doc.add_paragraph().add_run("TEST REPORT\n"), 28)

    table = doc.add_table(rows=2, cols=1, style="Table Grid")
    fPara = table.cell(0, 0).paragraphs[0]
    format_paragraph_spacing(fPara)
    fRun = fPara.add_run("Test Type: AS 4420.1 for Compliance to AS 2047")
    format_text(fRun, 12, bold_status=True, underline_status=True)
    format_text(table.cell(0, 0).paragraphs[0].add_run("          Report Issued: "), 12, bold_status=True)
    add_field(table.cell(0, 0).paragraphs[0])

    labels = ["Product Type", "Product Name", "Report Number", "Location", "Client", "Client Info", "Tested By", "Date Tested", "Checked By"]
    values = body_info + ["Toby Xie"]  # Appending name if not already included
    summary = "\n".join(f"{k}: {v}" for k, v in zip(labels, values))

    format_text(table.cell(1, 0).paragraphs[0].add_run(summary), 12, bold_status=True)


def _add_signature(doc):
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=1, style="Table Grid")
    sub = table.rows[0].cells[0].add_table(rows=1, cols=3)
    
    format_text(sub.cell(0, 0).paragraphs[0].add_run("Authorized Signatory: "), 12, bold_status=True)
    p = sub.cell(0, 1).paragraphs[0]
    p.add_run().add_picture(get_asset_path(SIGNATURE_IMAGE))
    format_text(p.add_run("\nToby Xie, "), 12)
    format_text(p.add_run("Research and Laboratory Manager, CiiLOCK Engineering\n"), 9, italicize=True)
    add_field(p)

    # Add footer block of text
    text_cell = table.rows[0].cells[0].add_table(rows=1, cols=1).cell(0, 0).paragraphs[0]
    footer = (
        "Accredited for compliance with ISO/IEC 17025 – Testing \n\nNATA is a signatory to the ILAC Mutual Recognition Arrangement for the mutual recognition of the equivalence of testing, medical testing, calibration, inspection, proficiency testing scheme providers and reference materials producers reports and certificates\n\nThe results apply to sample as received."
    )
    format_text(text_cell.add_run(footer), 12, bold_status=True)


def _add_footer(doc, report_number):
    section = doc.sections[0]
    footer = section.footer

    table = footer.add_table(rows=2, cols=1, width=section.page_width)
    inner = table.cell(0, 0).add_table(rows=1, cols=2)

    format_text(inner.cell(0, 0).paragraphs[0].add_run(report_number), 12)
    p = inner.cell(0, 1).paragraphs[0]
    format_text(p.add_run("Page "), 12)
    add_page_number(p.add_run(), "PAGE")
    format_text(p.add_run(" of "), 12)
    add_page_number(p.add_run(), "NUMPAGES")
    inner.cell(0, 1).paragraphs[0].alignment = 2

    disclaimer = (
        "Accredited for compliance with ISO/IEC 17025 – Testing. This document shall not be reproduced, except in full. Sample information was supplied by the client, and no verification of actual construction details or sampling of production stock could be performed. CiiLOCK Engineering accepts no liability for claims of losses, expenses, damages, and costs arising as result of the use of the products referred to this report."
    )
    format_text(table.cell(1, 0).paragraphs[0].add_run(disclaimer), 8)