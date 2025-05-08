from docx import Document
from src.utils.Formatter import *

# Run all functions to create the appendix
def start_all(doc: Document):
    appendix_start(doc)
    appendix_product_rating(doc)

# Specimen Details
def appendix_start(doc: Document):
    """
    Creates the start of the appendix section
    """
    appSpecimen = doc.add_paragraph().add_run("Appendix. Specimen details\n")
    format_text(appSpecimen, 20, bold_status=True, underline_status=True)
    appTestSetup = doc.add_paragraph().add_run("Appendix. Test Setup\n")
    format_text(appTestSetup, 20, bold_status=True, underline_status=True)
    appDisplace = doc.add_paragraph().add_run("Appendix. Displacement sensor reading for each pressure increment.\n\n\n")
    format_text(appDisplace, 20, bold_status=True, underline_status=True)
    listRun1 = doc.add_paragraph("", style='List Bullet').add_run("The term “Span*” denotes the length measured from the end displacement sensor probes.")
    format_text(listRun1, 7)
    listRun2 = doc.add_paragraph("", style='List Bullet').add_run("Unless otherwise stated, the default configuration places probes at both ends, with an additional sensor positioned at the midpoint.")
    format_text(listRun2, 7)
    listRun3 = doc.add_paragraph("", style='List Bullet').add_run("The initial sensor reading at pre-zero pressure indicates the start of initial compression reading of the sensor probe, without indicating any displacement of the specimen.")
    format_text(listRun3, 7)
    listRun4 = doc.add_paragraph("", style='List Bullet').add_run("Displacement readings have been taken on the test specimen at nominated locations represent the maximum expected structural movement. The location of all displacement measuring devices were agreed between the client and the CiiLOCK Engineering.")
    format_text(listRun4, 7)

    appProduct = doc.add_paragraph().add_run("Appendix. Product Rating from AS2047")
    format_text(appProduct, 20, bold_status=True, underline_status=True)

# Appendix. Product Rating from AS2047
def appendix_product_rating(doc):
    """
    Runs the code sections for product ratings
    """
    serviceability_rating(doc)
    maximum_air_infiltration(doc)
    water_penetration_resistance(doc)
    ultimate_strength(doc)
    operational_force(doc)

# Serviceability Rating
def serviceability_rating(doc):
    """
    Creates the serviceability rating section of the appendix
    """
    fPara = doc.add_paragraph()
    format_paragraph_spacing(fPara, 1, 0)
    format_text(fPara.add_run("Serviceability Rating"), 12, underline_status=True)
    serTable = doc.add_table(rows=11, cols=3, style="Table Grid")
    serTable.auto_fit = False
    
    titles = ["Serviceability Rating", "Test Pressure (Pa)", "Corner Window Test Pressure (Pa)"]
    ratings = [("N1", 400, 600),
    ("N2", 400, 600),
    ("N3", 600, 800),
    ("N4", 800, 1200),
    ("N5", 1200, 1800),
    ("N6", 1600, 2500),
    ("C1", 600, 900),
    ("C2", 800, 1200),
    ("C3", 1200, 1800),
    ("C4", 1600, 2500) ]

    # Add titles
    for idx in range(3):
        cell = serTable.cell(0, idx)
        cell.paragraphs[0].add_run(titles[idx])
        format_text(cell.paragraphs[0].runs[0], 12, bold_status=True)
        cell.paragraphs[0].alignment = 1
    
    # Add content for serviceability ratings
    for idx in range(10):
        rating, gen, corn = ratings[idx]
        ratingCell = serTable.cell(1+idx, 0)
        genCell = serTable.cell(1+idx, 1)
        cornCell = serTable.cell(1+idx, 2)
        
        ratingCell.paragraphs[0].add_run(rating)
        genCell.paragraphs[0].add_run(str(gen))
        cornCell.paragraphs[0].add_run(str(corn))

        for run in [ratingCell, genCell, cornCell]:
            format_text(run.paragraphs[0].runs[0], 12)
            run.paragraphs[0].alignment = 1
    
    # format table
    for cell in serTable.columns[0].cells:
        format_table_width(cell, 2)
    for idx in range(2):
        for cell in serTable.columns[idx+1].cells:
            format_table_width(cell, 1.5)

# Maximum Air Infiltration
def maximum_air_infiltration(doc):
    """
    creates the section for air infiltration
    """
    fPara = doc.add_paragraph()
    format_paragraph_spacing(fPara, 1, 0)
    format_text(fPara.add_run("\nMaximum Air Infiltration"), 12, underline_status=True)
    aITable = doc.add_table(rows=3, cols=3, style="Table Grid")
    aITable.auto_fit = False

    titles = [("Air infiltration level", "Low", "High"), 
              ("Pressure direction", "Positive and Negative", "Positive")]
    col = 0
    idx = 0
    # Add titles
    for idx2 in range(2):
        v1, v2, v3 = titles[idx2]
        cell1 = aITable.cell(idx, col)
        cell2 = aITable.cell(idx + 1, col)
        cell3 = aITable.cell(idx + 2, col)

        cell1.paragraphs[0].add_run(v1)
        cell2.paragraphs[0].add_run(v2)
        cell3.paragraphs[0].add_run(v3)

        for cell in [cell1, cell2, cell3]:
            format_text(cell.paragraphs[0].runs[0], 12)
            cell.paragraphs[0].alignment = 1
        col += 1

    format_text(aITable.cell(0, 0).paragraphs[0].runs[0], 12, bold_status=True)
    format_text(aITable.cell(0, 1).paragraphs[0].runs[0], 12, bold_status=True)

    # Add final column
    cell = aITable.cell(0, 2)
    cell.paragraphs[0].add_run("Measured Flow (L/sm")
    cell.paragraphs[0].add_run("2").font.superscript = True
    cell.paragraphs[0].add_run(", Test pressure, 75 Pa")
    cell.paragraphs[0].alignment = 1  # Center alignment
    for run in cell.paragraphs[0].runs:
        format_text(run, 12, bold_status=True)
    cell.paragraphs[0].alignment = 1
    
    format_text(aITable.cell(1, 2).paragraphs[0].add_run("1"), 12)
    format_text(aITable.cell(2, 2).paragraphs[0].add_run("5"), 12)
    aITable.cell(1, 2).paragraphs[0].alignment = 1
    aITable.cell(2, 2).paragraphs[0].alignment = 1

    for idx in range(3):
        for cell in aITable.columns[idx].cells:
            format_table_width(cell, 1.5)

# Water Peneration Resistance
def water_penetration_resistance(doc):
    """
    creates the table for water penetration resistance
    """
    fPara = doc.add_paragraph()
    format_paragraph_spacing(fPara, 1)
    format_text(fPara.add_run("\nWater Penetration Resistance"), 12, underline_status=True)
    
    titles = ["Window Ratings", "Non-Exposed", "Exposed"]
    waterRatings = [
    ("N1, N2", 150, 200),
    ("N3, C1", 150, 300),
    ("N4, C2", 200, 300),
    ("N5, C3", 300, 450),
    ("N6, C4", 450, 600)]

    waterTable = doc.add_table(rows=7, cols=3, style="Table Grid")
    waterTable.auto_fit = False
    # merge top cell
    waterTable.cell(0, 0).merge(waterTable.cell(0, 2))
    run = waterTable.cell(0, 0).paragraphs[0].add_run("Water Penetration Resistance Test Pressure (Pa)")
    format_text(run, 12, bold_status=True)

    for idx in range(3):
        cell = waterTable.cell(1, idx)
        cell.paragraphs[0].add_run(titles[idx])
        format_text(cell.paragraphs[0].runs[0], 12, bold_status=True)

    
    for idx in range(5):
        rating, nExp, exp = waterRatings[idx]
        ratingCell = waterTable.cell(2+idx, 0)
        genCell = waterTable.cell(2+idx, 1)
        cornCell = waterTable.cell(2+idx, 2)
        
        ratingRun = ratingCell.paragraphs[0].add_run(rating)
        nExpRun = genCell.paragraphs[0].add_run(str(nExp))
        expRun = cornCell.paragraphs[0].add_run(str(exp))

        for run in [ratingRun, nExpRun, expRun]:
            format_text(run, 12)
    
    # format 
    for idx in range(3):
        for cell in waterTable.columns[idx].cells:
            format_table_width(cell, 1.5)
            cell.paragraphs[0].alignment = 1

# Ultimate Strength
def ultimate_strength(doc):
    """
    creates the section for ultimate strength in the appendix
    """
    fPara = doc.add_paragraph("\n")
    format_paragraph_spacing(fPara, 1)
    format_text(fPara.add_run("Ultimate Strength"), 12, underline_status=True)
    ultStrRatings = [
    ("N1", 600, 900),
    ("N2", 900, 1300),
    ("N3", 1400, 2000),
    ("N4", 2000, 3000),
    ("N5", 3000, 4500),
    ("N6", 4000, 6000),
    ("C1", 1800, 2700),
    ("C2", 2700, 4000),
    ("C3", 4000, 5900),
    ("C4", 5300, 8000)]
    ultStrTitles = ["Window Rating", "General Window Test Pressure (Pa)", "Corner Window Test Pressure (Pa)"]

    # Add table
    ultStrTable = doc.add_table(rows=1+len(ultStrRatings), cols=len(ultStrTitles), style="Table Grid")
    ultStrTable.auto_fit = False

    # Add title
    for idx in range(len(ultStrTitles)):
        cell = ultStrTable.cell(0, idx)
        cell.paragraphs[0].add_run(ultStrTitles[idx])
        format_text(cell.paragraphs[0].runs[0], 12, bold_status=True)

    # Add ratings
    for idx in range(len(ultStrRatings)):
        rating, gen, corn = ultStrRatings[idx]
        ratingCell = ultStrTable.cell(1+idx, 0)
        genCell = ultStrTable.cell(1+idx, 1)
        cornCell = ultStrTable.cell(1+idx, 2)
        
        ratingRun = ratingCell.paragraphs[0].add_run(rating)
        genRun = genCell.paragraphs[0].add_run(str(gen))
        cornRun = cornCell.paragraphs[0].add_run(str(corn))

        for run in [ratingRun, genRun, cornRun]:
            format_text(run, 12)
        
    # format table width
    for idx in range(3):
        for cell in ultStrTable.columns[idx].cells:
            cell.paragraphs[0].alignment = 1
            format_table_width(cell, 1.75)

# Allowable Operational Force
def operational_force(doc):
    """
    Create section for operational force
    """
    fPara = doc.add_paragraph("\n")
    format_paragraph_spacing(fPara, 1)
    format_text(fPara.add_run("Allowable Operational Force"), 12, underline_status=True)
    oFTitles = ["Product Type", "Force (N) Initiate", "Force (N) Maintain"]
    oFRatings = [("Horizontal Sliding Window", 110, 90),
    ("Vertical Sliding Window", 200, 160),
    ("Horizontal Sliding Door", 180, 110),
    ("Swinging Door", 60, 20),
    ("Projecting Sash (no operator)", 160, 80)]
    
    oFTable = doc.add_table(rows=len(oFRatings) + 1, cols=len(oFTitles), style="Table Grid")
    oFTable.auto_fit = False

    for idx in range(len(oFTitles)):
        cell = oFTable.cell(0, idx)
        cell.paragraphs[0].add_run(oFTitles[idx])
        format_text(cell.paragraphs[0].runs[0], 12, bold_status=True)
        cell.paragraphs[0].alignment = 1
    
    # Add ratings
    for idx in range(len(oFRatings)):
        rating, initVal, maintVal = oFRatings[idx]
        desc = oFTable.cell(1+idx, 0).paragraphs[0].add_run(rating)
        init = oFTable.cell(1+idx, 1).paragraphs[0].add_run(str(initVal))
        oFTable.cell(1+idx, 1).paragraphs[0].alignment = 1
        maint = oFTable.cell(1+idx, 2).paragraphs[0].add_run(str(maintVal))
        oFTable.cell(1+idx, 2).paragraphs[0].alignment = 1
        

        for run in [desc, init, maint]:
            format_text(run, 12)
    
    for idx in range(3):
        for cell in oFTable.columns[idx].cells:
            format_table_width(cell, 1.5)
    
    