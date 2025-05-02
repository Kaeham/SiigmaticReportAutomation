from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from src.DataReading import DataReader
import os
import sys
from NRatingFinder import *
from FormattingFunctions import *
from Appendix import start_all
from tkinter import Tk, filedialog
import datetime

class ReportGenerator():
    def __init__(self):
        """
        Inputs:
            - file: the excel file containing all the data
        """
        Tk().withdraw()
        file = filedialog.askopenfilename(title="Select Excel File")
        self.reader = DataReader(file)
        self.data = self.reader.get_data()

        # Data variables
        self.bodyData = self.data[0] # 1 row
        self.deflectionData = self.data[1] # requires len*2 rows
        self.oFData = self.data[2] # requires len rows
        self.aIData = self.data[3] # requires 2 rows
        self.waterData = self.data[4] # requires 1 row
        self.ultimateData = self.data[5] # requires 1 row
        self.deflectionVal = self.data[6] #
        self.filename = self.data[7] #

        self.sashCount = max(2, len(self.deflectionData))
        if self.oFData != None:
            self.oFSashCount = max(2, len(self.oFData)) 
        else:
            self.oFSashCount = 2

        # Define index variables for each variable
        self.summaryRows = 6 + self.sashCount*2 + self.oFSashCount
        self.titleRowIdx = 0
        self.serRowIdx = self.titleRowIdx + 1
        self.defRowIdx = self.serRowIdx + 1
        self.oFRowIdx = self.defRowIdx + self.sashCount*2
        self.aIRowIdx = self.oFRowIdx + self.oFSashCount
        self.waterRowIdx = self.aIRowIdx + 2
        self.ultimateRowIdx = self.waterRowIdx + 1
        self.errors = []

        self.document = Document()
        self.title_page(self.document)
        self.summary_results(self.document)
        self.test_results(self.document)
        start_all(self.document)

    def get_asset_path(self, filename):
        """ 
        Get the correct path for assets (like images) when running as .exe
        :Inputs:
            - filename: the name of the file to get the path for
        """
        if getattr(sys, 'frozen', False):  # Running as .exe
            base_path = sys._MEIPASS  # Temporary folder where PyInstaller extracts files
        else:
            base_path = os.path.dirname(os.path.abspath(__file__))  # Running as .py script

        return os.path.join(base_path, filename)

    def header(self, doc: Document):
        """
        creates and formats the header for the document
        Inputs:
            - doc: document that is being generated
        """
        section = doc.sections[0]
        header = section.header
        ciilockPath = self.get_asset_path("img/ciilock.jpg")
        nataPath = self.get_asset_path("img/natalogo.jpg")

        # add table
        headerTable = header.add_table(rows=1, cols=3, width=section.page_width+section.right_margin)
        headerTable.autofit = False # ensure the last column can go to the edge of the page
        headerTable.cell(0, 0).paragraphs[0].add_run().add_picture(ciilockPath) #
        headerTableText = headerTable.cell(0, 1).paragraphs[0] #
        
        # format and add header contents
        format_text(headerTableText.add_run("CiiLOCK Engineering\n"), 20, color=RGBColor(54, 95, 145))
        format_text(headerTableText.add_run("ABN: 90 112 576 815     PH.\n"), 12)
        format_text(headerTableText.add_run("18 Technology Circuit Factory 2, Hallam Vic 3803"), 12)
        headerTable.cell(0, 2).paragraphs[0].alignment = 2
        format_table_width(headerTable.cell(0, 0), 1)
        format_table_width(headerTable.cell(0, 1), 4)
        format_table_width(headerTable.cell(0, 2), 2)
        format_table_height(headerTable.rows[0], 1)
        section.header_distance = Pt(0)
        headerTable.cell(0, 2).paragraphs[0].add_run().add_picture(nataPath)

    def body(self, doc: Document):
        """
        Adds the information present in the body of the title page
        Inputs:
            - doc: document that is being generated
        """
        format_text(doc.add_paragraph().add_run(""), 20)
        format_text(doc.add_paragraph().add_run("TEST REPORT\n", 0), 28)
        # Table with details
        tableOne = doc.add_table(rows=2, cols=1, style="Table Grid")
        fPara = tableOne.cell(0, 0).paragraphs[0]
        format_paragraph_spacing(fPara)
        fRun = fPara.add_run("Test Type: AS 4420.1 for Compliance to AS 2047")
        format_text(fRun, 12, bold_status=True, underline_status=True)
        sRun = fPara.add_run("          Report Issued: ")
        format_text(sRun, 12, bold_status=True)
        add_field(fPara)

        inputs = ["Product Type: ", "Product Name: ", "Report Number: ", "Location: ", "Client: ", "Client Information: ",
                "Tested By: ", "Date Tested: ", "Report Checked By: "]
        res = ""

        self.bodyData.append("Toby Xie")
        for idx in range(len(inputs)):
            res += inputs[idx] + str(self.bodyData[idx]) + "\n"

        format_text(tableOne.cell(1, 0).paragraphs[0].add_run(res), 12, bold_status=True)
        
        # Signature Section
        doc.add_paragraph()
        tableTwo = doc.add_table(rows=1, cols=1, style="Table Grid")
        tableTwoSubTableOne = tableTwo.rows[0].cells[0].add_table(rows=1, cols=3)
        tableTwoSubTableTwo = tableTwo.rows[0].cells[0].add_table(rows=1, cols=1)
        signatureTop = tableTwoSubTableOne.cell(0, 0).paragraphs[0].add_run("Authorized Signatory: ")
        format_text(signatureTop, 12, bold_status=True)
        signaturePath = self.get_asset_path("img/signature.png")  # Get correct image path
        
        imageCell = tableTwoSubTableOne.cell(0, 1)
        imageCell.paragraphs[0].add_run().add_picture(signaturePath)
        imageCell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Signature Credentials
        imageCellText = imageCell.paragraphs[0]
        signatureBodyOne = imageCellText.add_run("\nToby Xie, ")
        format_text(signatureBodyOne, 12)
        signatureBodyTwo = imageCellText.add_run("Research and Laboratory Manager, CiiLOCK Engineering\n")
        format_text(signatureBodyTwo, 9, italicize=True)
        add_field(imageCellText)
        format_table_height(tableTwoSubTableOne.rows[0], 1.5)
        format_table_width(tableTwoSubTableOne.rows[0].cells[0], 1.5)
        format_table_width(tableTwoSubTableOne.rows[0].cells[1], 4.5)
        format_table_width(tableTwoSubTableOne.rows[0].cells[2], 1.5)

        # Signature Text
        textCell = tableTwoSubTableTwo.cell(0, 0).paragraphs[0]
        signatureFinal = textCell.add_run("Accredited for compliance with ISO/IEC 17025 – Testing \n\nNATA is a signatory to the ILAC Mutual Recognition Arrangement for the mutual recognition of the equivalence of testing, medical testing, calibration, inspection, proficiency testing scheme providers and reference materials producers reports and certificates\n\nThe results apply to sample as received.")
        format_text(signatureFinal, 12, bold_status=True)

    def footer(self, doc: Document):
        """
        creates and formats the footer for the document
        doc: document that is being generated
        """
        section = doc.sections[0]
        section.footer_distance = Pt(0)

        footer = section.footer
        # Table containing the tables for the sub section
        # Row 1 contains test report number and page number
        # Row 2 contains the text
        containerTable = footer.add_table(rows=2, cols=1, width=section.page_width)
        subTable = containerTable.cell(0, 0).add_table(rows=1, cols=2)
        pagNum = "Page "
        reportNum = str(self.bodyData[2])
        subTable.cell(0, 0).paragraphs[0].add_run(reportNum)
        format_text(subTable.cell(0, 0).paragraphs[0].runs[0], 12)
        run = subTable.cell(0, 1).paragraphs[0].add_run(pagNum)
        format_text(run, 12)
        add_page_number(run, 'PAGE')
        run = subTable.cell(0, 1).paragraphs[0].add_run(" of ")
        add_page_number(run, 'NUMPAGES')
        format_text(run, 12)

        subTable.cell(0, 1).paragraphs[0].alignment = 2

        res = "Accredited for compliance with ISO/IEC 17025 – Testing. This document shall not be reproduced, except in full. Sample information was supplied by the client, and no verification of actual construction details or sampling of production stock could be performed. CiiLOCK Engineering accepts no liability for claims of losses, expenses, damages, and costs arising as result of the use of the products referred to this report."
        footerText = containerTable.cell(1, 0).paragraphs[0].add_run(res)
        format_text(footerText, 8)

    def title_page(self, doc: Document):
        """
        Function which creates the title page
        Inputs:
            - doc: document that is being generated
        """
        self.header(doc)
        self.body(doc)
        self.footer(doc)
        doc.add_page_break()

    def fill_in_summary(self, table):
        """
        Fills in the summary page on the document
        Inputs:
            - table: the table object on the summary page    
        """
        serRatingGen, serRatingCorn = n_rating_serviceability(self.deflectionVal)
        aIRating = n_rating_air_results(self.aIData[0], self.aIData[1])
        waterRatingExposed, waterRatingNonExposed = n_rating_water(self.waterData[0])
        genRatingUST, cornRatingUST = n_rating_ust(min(self.ultimateData[0], self.ultimateData[1]))

        # Serviceability Design Wind Pressure
        table.cell(self.serRowIdx, 1).text = "±" + str(self.deflectionVal) + " Pa"
        if serRatingGen and serRatingCorn:
            table.cell(self.serRowIdx, 3).text = serRatingGen + "\n" + serRatingCorn
        table.cell(self.serRowIdx, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Deflection/Span Ratio
        member = 1
        plus = True
        for x in range(self.sashCount*2):
            res = f"Vertical Sash on Panel (V0{member}){'+' if plus else '-'}" 
            table.cell(self.defRowIdx + x, 1).text = res
            plus = not plus
            if x % 2 == 1:
                member += 1
        
        idx = self.defRowIdx
        print(self.deflectionData)
        for member in self.deflectionData:
            member = member[1:]
            for span in member:
                table.cell(idx, 2).text = "Span/" + str(span)
                idx += 1

        # Operation Force
        panel = 0
        for idx in range(0, self.oFSashCount, 2):
            panel += 1
            resOpen = f"Panel {panel} Open Initiate / Maintain"
            resClose = f"Panel {panel} Close Initiate / Maintain"
            table.cell(self.oFRowIdx + idx, 1).text = resOpen
            table.cell(self.oFRowIdx + idx + 1, 1).text = resClose

        if len(self.oFData) >= self.oFSashCount:
            for idx in range(0, len(self.oFData)):
                if len(self.oFData) > idx and len(self.oFData[idx]) == 2:
                    init, maint = self.oFData[idx]
                    # print(self.oFData[idx])
                else:
                    init, maint = "N/A", "N/A"
                table.cell(self.oFRowIdx + idx, 2).text = str(init) + " / " + str(maint)
        else:
            self.errors.append("Error: Operating Force Data not recorded properly")

        # Air Infiltration
        if len(self.aIData) >= 2:
            table.cell(self.aIRowIdx, 1).text = "At +75Pa"
            table.cell(self.aIRowIdx + 1, 1).text = "At -75Pa"
            table.cell(self.aIRowIdx, 2).text = self.aIData[0] + "/sm²"
            if self.aIData[1] == 0:
                table.cell(self.aIRowIdx + 1, 2).text = "NR"
            else:
                table.cell(self.aIRowIdx + 1, 2).text = self.aIData[0] + "/sm²"
            if aIRating:
                table.cell(self.aIRowIdx, 3).text = aIRating
        else:
            self.errors.append("Error: Air Infiltration Data not recorded properly")

        # Water Test
        if len(self.waterData) >= 1:
            table.cell(self.waterRowIdx, 1).text = str(self.waterData[0]) + "Pa"
            table.cell(self.waterRowIdx, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            rating = ""
            if waterRatingExposed:
                rating += waterRatingExposed + "\n"
            if waterRatingNonExposed:
                rating += waterRatingNonExposed
            table.cell(self.waterRowIdx, 3).text = rating

        else:
            self.errors.append("Error: Water Penetration not recorded properly")

        # Ultimate Strength Test
        if self.ultimateData:
            posUTS, negUTS = self.ultimateData[0], self.ultimateData[1]
            table.cell(self.ultimateRowIdx, 1).text = f"+{str(posUTS)}Pa\n-{str(negUTS)}Pa"
            table.cell(self.ultimateRowIdx, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if genRatingUST and cornRatingUST:
                table.cell(self.ultimateRowIdx, 3).text = genRatingUST + "\n" + cornRatingUST
        else:
            self.errors.append("Error: Ultimate Data not recorded properly")

    def summary_results(self, doc: Document):
        """
        Creates and formats the summary page
        Inputs:
            - doc: document that is being generated
        """
        summary = doc.add_paragraph(style="Heading 1").add_run("Summary")
        format_text(summary, 20, bold_status=True, underline_status=True)

        table = doc.add_table(rows=self.summaryRows, cols=5, style="Table Grid")
        # work through the table column by column
        table.cell(self.defRowIdx, 0).merge(table.cell(self.oFRowIdx - 1, 0)) # column 1 merging
        table.cell(self.oFRowIdx, 0).merge(table.cell(self.aIRowIdx - 1, 0))
        table.cell(self.aIRowIdx, 0).merge(table.cell(self.waterRowIdx - 1, 0)) 
        table.cell(self.titleRowIdx, 0).paragraphs[0].add_run("Test Description").bold = True # col 1 title 
        table.cell(self.serRowIdx, 0).text = "Serviceability Design Wind Pressure"
        table.cell(self.defRowIdx, 0).text = "Deflection/Span Ratio Structural member 1"
        table.cell(self.oFRowIdx, 0).text = "Operational Force"
        table.cell(self.aIRowIdx, 0).text = "Air infiltration at +- 75Pa"
        table.cell(self.waterRowIdx, 0).text = "Water Penetration"
        table.cell(self.ultimateRowIdx, 0).text = "Ultimate Strength Test Pressure"
        for row in table.rows:
            format_table_width(row.cells[0], 1.5)
            format_table_width(row.cells[1], 2)
            # format_table_width(row.cells[2], 1)
            format_table_width(row.cells[3], 1.5)
            format_table_width(row.cells[4], 1)

        table.cell(self.titleRowIdx, 1).merge(table.cell(self.titleRowIdx, 2)) # column 2 & 3
        table.cell(self.serRowIdx, 1).merge(table.cell(self.serRowIdx, 2))
        table.cell(self.waterRowIdx, 1).merge(table.cell(self.waterRowIdx, 2))
        table.cell(self.ultimateRowIdx, 1).merge(table.cell(self.ultimateRowIdx, 2))
        table.cell(self.titleRowIdx, 1).paragraphs[0].add_run("Test Results").bold = True
        table.cell(self.aIRowIdx, 1).paragraphs[0].add_run("+ 75")
        table.cell(self.aIRowIdx + 1, 1).paragraphs[0].add_run("- 75")
        table.cell(self.serRowIdx, 3).merge(table.cell(self.oFRowIdx - 1, 3)) # column 4
        table.cell(self.oFRowIdx, 3).merge(table.cell(self.aIRowIdx - 1, 3))
        table.cell(self.aIRowIdx, 3).merge(table.cell(self.waterRowIdx - 1, 3))
        table.cell(self.titleRowIdx, 3).paragraphs[0].add_run("Rating").bold = True
        table.cell(self.serRowIdx, 4).merge(table.cell(self.oFRowIdx - 1, 4)) # column 5
        table.cell(self.oFRowIdx, 4).merge(table.cell(self.aIRowIdx - 1, 4))
        table.cell(self.aIRowIdx, 4).merge(table.cell(self.waterRowIdx-1, 4))
        table.cell(self.titleRowIdx, 4).paragraphs[0].add_run("Verdict").bold = True

        fPara = doc.add_paragraph()
        note = fPara.add_run("NR: measurement not required, herinafter.\n\n")
        format_text(note, 12)

        self.fill_in_summary(table)
        format_summary_page(table)
        insert_line_break(fPara)
        doc.add_page_break()

    def test_results(self, doc: Document):
        """
        runs functions for test results, which create and format the page for each test
        Inputs:
            - doc: document that is being generated
        """
        format_text(doc.add_paragraph(style="Heading 1").add_run("Test Results\n"), 20, bold_status=True, underline_status=True)
        self.deflection(doc)
        self.operating_force_results(doc)
        self.air_results(doc)
        self.water_results(doc)
        self.ultimate_results(doc)

    def deflection(self, doc: Document):
        """
        creates and formats the deflection test page
        Inputs:
            - doc: document that is being generated
        """
        # Add Title "Deflection"
        dPara = doc.add_paragraph()
        dTitle = dPara.add_run("Deflection\n")
        format_text(dTitle, 14, bold_status=True, underline_status=True)
        # Add Method Section
        format_text(dPara.add_run("Method\n"), 12, bold_status=True)
        format_text(dPara.add_run("The test specimen was tested in accordance with AS4420.1 Clause 3.\n"), 12)

        # Add Results Heading
        format_text(dPara.add_run("\nResults"), 12, bold_status=True)
        format_paragraph_spacing(dPara)

        # Define Sashes
        for sash in self.deflectionData:
            resultPara = doc.add_paragraph()
            format_paragraph_spacing(resultPara)
            name, pos, neg = sash
            format_text(resultPara.add_run(name.capitalize()), 12)

            # Create Table (3 rows, 3 columns)
            table = doc.add_table(rows=3, cols=3, style="Table Grid")
            # Add Headers
            table.cell(0, 0).paragraphs[0].add_run("Pressure Differential").bold = True
            table.cell(0, 1).paragraphs[0].add_run("Pressure Achieved (Pa)").bold = True
            table.cell(0, 2).paragraphs[0].add_run("Resultant Span Ratio").bold = True

            # Add Row 1 (Positive Pressure)
            table.cell(1, 0).paragraphs[0].add_run("Positive").bold = True
            table.cell(1, 1).paragraphs[0].add_run(str(self.deflectionVal))
            table.cell(1, 2).paragraphs[0].add_run(f"Span/{pos}")

            # Add Row 2 (Negative Pressure)
            table.cell(2, 0).paragraphs[0].add_run("Negative").bold = True
            table.cell(2, 1).paragraphs[0].add_run(f"-{str(self.deflectionVal)}")
            table.cell(2, 2).paragraphs[0].add_run(f"Span/{neg}")
            
            # Formatting First Row
            for cell in table.rows[0].cells:
                format_text(cell.paragraphs[0].runs[0], 12, bold_status=True)
            
            # Formatting First Column
            for row in table.rows:
                format_text(row.cells[0].paragraphs[0].runs[0], 12, bold_status=True)
            
            for row in table.rows[1:]:
                for cell in row.cells[1:]:    
                    format_text(cell.paragraphs[0].runs[0], 12)
                    format_cell_color(cell, cellColor)

            doc.add_paragraph()
            
        format_text(doc.add_paragraph().add_run("(See sash numbering in Appendix: Elevation drawing)"), 12, italicize=True)
        
        # Add Compliance Section
        finalPara = doc.add_paragraph()
        format_text(finalPara.add_run("Compliance\n"), 12, bold_status=True)
        format_text(finalPara.add_run("The test specimen satisfied the deflection requirements of AS2047 Clause 2.3.1.3 at the specified pressure.\n\n"), 12)
        format_text(finalPara.add_run("Comment: Nil"), 12)
        insert_line_break(finalPara)
        doc.add_page_break()

    def operating_force_results(self, doc: Document):
        """
        creates and formats the operating force results
        Inputs:
            - doc: document that is being generated
        """
        # Add title
        oFPara = doc.add_paragraph()
        format_text(oFPara.add_run("\nOperating Force\n"), 14, bold_status=True, underline_status=True)

        # Add "Method" section
        format_text(oFPara.add_run("Method\n"), 12, bold_status=True)
        format_text(oFPara.add_run("The test specimen was tested in accordance with AS4420.1 Clause 4.\n"), 12)

        # Add "Results" section
        format_text(oFPara.add_run("\nResults"), 12, bold_status=True)

        panel = 0
        for idx in range(0, len(self.oFData), 2):
            panel += 1
            init, maint = self.oFData[idx]
            panelPara = doc.add_paragraph()
            format_paragraph_spacing(panelPara)
            format_text(panelPara.add_run(f"Panel {panel}"), 12)
            
            # Create a table with 3 rows and 3 columns
            table = doc.add_table(rows=3, cols=3, style="Table Grid")

            # Fill table headers
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = ""
            hdr_cells[1].text = "Initiating Force"
            format_text(hdr_cells[1].paragraphs[0].runs[0], 12, bold_status=True)
            hdr_cells[2].text = "Maintaining Force"
            format_text(hdr_cells[2].paragraphs[0].runs[0], 12, bold_status=True)

            # First row (Open Force)
            row_cells = table.rows[1].cells
            row_cells[0].text = "Opening"
            format_text(row_cells[0].paragraphs[0].runs[0], 12, bold_status=True)
            row_cells[1].text = f"{init}N"
            format_cell_color(row_cells[1], cellColor)
            row_cells[2].text = f"{maint}N"
            format_cell_color(row_cells[2], cellColor)

            init, maint = self.oFData[idx+1]
            # Second row (Close Force)
            row_cells = table.rows[2].cells
            row_cells[0].text = "Closing"
            format_text(row_cells[0].paragraphs[0].runs[0], 12, bold_status=True)
            row_cells[1].text = f"{init}N"
            format_cell_color(row_cells[1], cellColor)
            row_cells[2].text = f"{maint}N"
            format_cell_color(row_cells[2], cellColor)

        # Add "Compliance" section
        endPara = doc.add_paragraph()
        format_text(endPara.add_run("\nCompliance\n"), 12, bold_status=True)
        format_text(endPara.add_run("The test specimen satisfied the requirements of AS2047 Clause 2.3.1.4\n"), 12)
        
        # Add "Comment" section
        format_text(endPara.add_run("\nComment: Nil\n"), 12)
        insert_line_break(endPara)
        doc.add_page_break()

    def air_results(self, doc: Document):
        """
        creates air results section in the document
        :Inputs:
            - doc: document that is being generated
        air infiltration data is in a list of tuples, where each tuple
        corresponds to a test. the 0th index and 1st index in tuple contains
        positive and negative air infiltration values respectively
        """
        # Add Title "Air Infiltration"
        titlePara = doc.add_paragraph()
        format_text(titlePara.add_run("\nAir Infiltration\n"), 14, bold_status=True, underline_status=True)
        
        # Add Method Section
        format_text(titlePara.add_run("Method\n"), 12, bold_status=True)
        format_text(titlePara.add_run("The test specimen was tested in accordance with AS4420.1 Clause 5.\n"), 12)

        # Add Results Section
        format_text(titlePara.add_run("Results"), 12, bold_status=True)

        # Create Table (4 rows, 4 columns)
        table = doc.add_table(rows=4, cols=5, style="Table Grid")

        # Merge top row for title
        table.cell(0, 0).merge(table.cell(0, 4))
        table.cell(0, 0).paragraphs[0].add_run("Measured Flow (Ls")
        table.cell(0, 0).paragraphs[0].add_run("-1").font.superscript = True
        table.cell(0, 0).paragraphs[0].add_run("m")
        table.cell(0, 0).paragraphs[0].add_run("-2").font.superscript = True
        table.cell(0, 0).paragraphs[0].add_run(")")
        table.cell(0, 0).paragraphs[0].alignment = 1  # Center alignment
        table.cell(2, 1).merge(table.cell(3, 1))

        # Add Headers
        headers = ["Air Infiltration Level", "Test Pressure (Pa)", "", "Positive Area Flow Rate", "Negative Area Flow Rate"]
        for i, text in enumerate(headers):
            table.cell(1, i).paragraphs[0].add_run(text)
        
        table.cell(1, 2).paragraphs[0].add_run("Allowable Flow (Ls")
        table.cell(1, 2).paragraphs[0].add_run("-1").font.superscript = True
        table.cell(1, 2).paragraphs[0].add_run("m")
        table.cell(1, 2).paragraphs[0].add_run("-2").font.superscript = True
        table.cell(1, 2).paragraphs[0].add_run(")")

        posVal, negVal = self.aIData
        if negVal == 0:
            negVal = "N/A"

        # Add Data Rows
        data = [
            ["Low", f"\u00B175", "< 1.0", posVal, negVal],
            ["High", "", "< 5.0", posVal, "N/A"]
        ]


        for row_idx, row_data in enumerate(data, start=2):
            for col_idx, value in enumerate(row_data):
                if value != "":
                    table.cell(row_idx, col_idx).text = str(value)
        
        # first and second row
        for cell in table.rows[0].cells:
            for run in cell.paragraphs[0].runs:
                format_text(run, 12, bold_status=True)
        for cell in table.rows[1].cells:
            for run in cell.paragraphs[0].runs:
                format_text(run, 12, bold_status=True)

        # second and third row, first and second column
        for row in table.rows[1:3]:
            for cell in row.cells[:2]:
                format_text(cell.paragraphs[0].runs[0], 12)

        # second and third row, 3-5 column
        for row in table.rows[2:]:
            for cell in row.cells[2:]:
                format_text(cell.paragraphs[0].runs[0], 12)
                format_cell_color(cell, cellColor)


        finalPara = doc.add_paragraph()
        # Add Compliance Section
        if float(posVal) < 1.0 and float(negVal) < 1.0:
            res = "Low"
        elif float(posVal) < 5.0:
            res = "High"
        
        format_text(finalPara.add_run("\nCompliance\n"), 12, bold_status=True)
        format_text(finalPara.add_run(f"The test specimen satisfied the “{res}” Air Infiltration Level requirements of AS2047 Clause 2.3.1.5\n"), 12)
        # Add Comment Section
        format_text(finalPara.add_run("\nComment: Nil\n"), 12)
        insert_line_break(finalPara)
        doc.add_page_break()

    def water_results(self, doc: Document):
        """
        plot water values on the word document
        :Inputs:
            - doc: document that is being generated
        waterData is a tuple with form: (measured water value, modification details, comments about test)
        """
        # Add Title
        wPParagraph = doc.add_paragraph()
        format_text(wPParagraph.add_run("\nWater Penetration\n"), 14, bold_status=True, underline_status=True)
        format_text(wPParagraph.add_run("Method\n"), 12, bold_status=True)
        format_text(wPParagraph.add_run("The test specimen was tested in accordance with AS4420.1 Clause 6.\n"), 12)
        format_text(wPParagraph.add_run("\nResults"), 12, bold_status=True)

        # Create Table (2 rows, 3 columns)
        table = doc.add_table(rows=2, cols=2, style="Table Grid")
        # Add Headers
        format_text(table.cell(0, 0).paragraphs[0].add_run("Test Pressure (Pa)"), 12, bold_status=True)
        # table.cell(0, 1).paragraphs[0].add_run("Modifications")
        format_text(table.cell(0, 1).paragraphs[0].add_run("Observations"), 12, bold_status=True)

        # Add Data
        waterVal, waterComments = self.waterData
        # if modDetails in [None, ""]:
        #     modDetails = "N/A"
        table.cell(1, 0).text = str(waterVal)
        table.cell(1, 1).text = waterComments

        # table.cell(1, 2).text = modDetails
        format_table_width(table.cell(1, 0), 2)
        format_table_width(table.cell(1, 1), 10)
        
        for cell in table.rows[1].cells:
            format_text(cell.paragraphs[0].runs[0], 12)
        
        format_cell_color(table.cell(1, 1), cellColor)

        # End paragraphs
        endPara = doc.add_paragraph()
        format_text(endPara.add_run("\nCompliance\n"), 12, bold_status=True)
        format_text(endPara.add_run(f"The test specimen satisfied the water penetration requirements of AS2047 Clause 2.3.1.6 at the specified pressure of {waterVal}Pa\n"), 12)
        format_text(endPara.add_run("\nComment: Nil"), 12)

        # Add Page Break
        insert_line_break(endPara)
        doc.add_page_break()

    def ultimate_results(self, doc: Document):
        """
        creates and formats the ultimate results section
        Inputs:
            - doc: document that is being generated
        """
        # print(len(ultimateData) > 0, isinstance(ultimateData, tuple), len(ultimateData) == 3)
        if len(self.ultimateData) > 0 and isinstance(self.ultimateData, tuple) and len(self.ultimateData) == 3:
            # print(ultimateData)
            posUTS, negUTS, uTSData = self.ultimateData
        else:
            posUTS, negUTS, uTSData = "N/A", "N/A", ["N/A"] * 10  # Default values

        # Add Title
        ust_title = doc.add_paragraph()
        format_text(ust_title.add_run("\nUltimate Strength Test\n"), 14, bold_status=True, underline_status=True) # Title
        format_text(ust_title.add_run("Method\n"), 12, bold_status=True)
        format_text(ust_title.add_run("The test was tested in accordance with AS4420.1 Clause 7\n"), 12)

        # Add Results Section
        format_text(ust_title.add_run("Results"), 12, bold_status=True)

        # Pressure Table
        pressuresTable = doc.add_table(rows=2, cols=2, style="Table Grid")
        pressuresTable.cell(0, 0).paragraphs[0].add_run("Positive Test Pressure Pa (Negative Chamber Pressure)")
        pressuresTable.cell(0, 1).paragraphs[0].add_run(str(posUTS))
        pressuresTable.cell(1, 0).paragraphs[0].add_run("Negative Test Pressure Pa (Positive Chamber Pressure)")
        pressuresTable.cell(1, 1).paragraphs[0].add_run(str(negUTS))
        # Format Pressure Table
        for row in pressuresTable.rows:
            format_text(row.cells[0].paragraphs[0].runs[0], 10, bold_status=True)
            format_text(row.cells[1].paragraphs[0].runs[0], 10)
            format_cell_color(row.cells[1], cellColor)
        format_table_width(pressuresTable.cell(0, 0), 3.5)
        format_table_width(pressuresTable.cell(0, 1), 1.5)
        
        doc.add_paragraph("\n")
        # Observations Table (Positive & Negative Pressure)
        observationsTable = doc.add_table(rows=7, cols=3, style="Table Grid")

        # Merge top row to make a single header
        observationsTable.cell(0, 0).merge(observationsTable.cell(0, 2))
        observationsTable.cell(0, 0).paragraphs[0].add_run("Observations - Pressure Differential").bold = True
        observationsTable.cell(0, 0).paragraphs[0].alignment = 1  # Center align

        # Set column headers
        observationsTable.cell(1, 0).paragraphs[0].add_run("")
        observationsTable.cell(1, 1).paragraphs[0].add_run("Positive Pressure (Y/N)")
        observationsTable.cell(1, 2).paragraphs[0].add_run("Negative Pressure (Y/N)")

        # Add Observation Data

        obs_data = [
        "Failure or dislodgement of any glazing",
        "Dislodgement of a frame or any part of a frame",
        "Removal of a light, either with or without its framing sash, from a frame",
        "Loss of support of a frame, such as when it is unstable in its opening in the building structure",
        "Failure of any sash, locking device, fastener or supporting stay, allowing an opening light to open"
        ]
        for row_idx, obs in enumerate(obs_data, start=0):
            if row_idx < len(uTSData) and row_idx + 5 < len(uTSData):
                pos = uTSData[row_idx]
                neg = uTSData[row_idx + 5]
            else:
                pos, neg = "N/A", "N/A"
            observationsTable.cell(row_idx + 2, 0).text = obs
            observationsTable.cell(row_idx + 2, 1).text = str(pos)
            observationsTable.cell(row_idx + 2, 2).text = str(neg)

        # Format Table
        # row 1-2, all columns
        for row in observationsTable.rows[:2]:
            for cell in row.cells:
                format_text(cell.paragraphs[0].runs[0], 11, bold_status=True)
        format_table_width(observationsTable.cell(1, 0), 4.5)
        # row 3-7, 1 column
        for row in observationsTable.rows[2:]:
            format_text(row.cells[0].paragraphs[0].runs[0], 9)
        # row 3-7, column 2-3
        for row in observationsTable.rows[2:]:
            for cell in row.cells[1:]:
                format_text(cell.paragraphs[0].runs[0], 9)
                format_cell_color(cell, cellColor)
                cell.paragraphs[0].alignment = 1
        
        observationsTable.cell(1, 1).paragraphs[0].alignment = 1
        observationsTable.cell(1, 2).paragraphs[0].alignment = 1


        # Compliance Section
        finalPara = doc.add_paragraph()
        format_text(finalPara.add_run("\nCompliance\n"), 12, bold_status=True)
        format_text(finalPara.add_run("The test specimen satisfied the Ultimate Strength requirements of AS2047 Clause 2.3.1.7 at the specified pressures.\n"), 12)
        # Comment Section
        format_text(finalPara.add_run("\nComment: Nil"), 12)

        insert_line_break(finalPara)
        doc.add_page_break()

    # def final(self):        
    #     start_all(new)

    def get_report(self):
        """
        return the currently generated report
        """
        return self.document

    def print_errors(self):
        """
        print all errors if they exist, typically used when saving the report
        """
        if len(self.errors):
            timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
            filename = f"error_log_{timestamp}.txt"

            with open(filename, "w") as errorFile:
                errorFile.write("ERROR LOG\n\n")
                for error in self.errors:
                    errorFile.write(error + "\n")
            
            print(f"Errors detected! See log: {errorFile}")

    def save_report(self):
        """
        Save the report
        """
        Tk().withdraw()
        savePath = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")], initialfile=f"{self.filename}.docx", title="Save Report As")
        if savePath:
            self.get_report().save(savePath)
            print(f"File saved at: {savePath}")
        else:
            print("No file location selected. File not saved.")
        self.print_errors()